"""Export endpoint — convert markdown to PDF or DOCX."""

import io
import logging

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from api.models.artifacts import ExportRequest

logger = logging.getLogger(__name__)
router = APIRouter()


def _md_to_html(md_text: str, title: str) -> str:
    """Convert markdown to a styled HTML document."""
    import markdown

    body = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8"/>
<title>{title}</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 2cm; color: #1e293b; font-size: 11pt; line-height: 1.6; }}
  h1 {{ font-size: 20pt; border-bottom: 2px solid #f97316; padding-bottom: 6px; }}
  h2 {{ font-size: 16pt; margin-top: 18pt; }}
  h3 {{ font-size: 13pt; margin-top: 14pt; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12pt 0; font-size: 10pt; }}
  th, td {{ border: 1px solid #cbd5e1; padding: 6px 10px; text-align: left; }}
  th {{ background: #f1f5f9; font-weight: 600; }}
  code {{ background: #f1f5f9; padding: 2px 5px; border-radius: 3px; font-size: 10pt; }}
  pre {{ background: #f8fafc; border: 1px solid #e2e8f0; padding: 10px; border-radius: 6px; overflow-x: auto; }}
  blockquote {{ border-left: 3px solid #f97316; padding-left: 12px; color: #475569; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


@router.post("")
async def export_artifact(req: ExportRequest):
    """Export markdown content to PDF, DOCX or raw MD."""
    if req.format == "md":
        buf = io.BytesIO(req.content.encode("utf-8"))
        return StreamingResponse(
            buf,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{req.title}.md"'},
        )

    if req.format == "pdf":
        try:
            from weasyprint import HTML  # type: ignore[import-untyped]
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="weasyprint não instalado. Execute: pip install weasyprint",
            )

        html_str = _md_to_html(req.content, req.title)
        try:
            pdf_bytes = HTML(string=html_str).write_pdf()
        except Exception as e:
            logger.exception("PDF generation failed: %s", e)
            raise HTTPException(status_code=500, detail=f"Erro ao gerar PDF: {e}")

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{req.title}.pdf"'},
        )

    if req.format == "docx":
        try:
            from docx import Document  # type: ignore[import-untyped]
            from docx.shared import Pt
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="python-docx não instalado. Execute: pip install python-docx",
            )

        try:
            doc = Document()
            doc.add_heading(req.title, level=1)

            for line in req.content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("| "):
                    # Simplified table: just add as paragraph with monospace feel
                    p = doc.add_paragraph(stripped)
                    for run in p.runs:
                        run.font.size = Pt(9)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
        except Exception as e:
            logger.exception("DOCX generation failed: %s", e)
            raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {e}")

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{req.title}.docx"'},
        )

    raise HTTPException(status_code=400, detail=f"Formato não suportado: {req.format}")
